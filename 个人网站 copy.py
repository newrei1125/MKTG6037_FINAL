import streamlit as st
import pandas as pd
from PIL import Image
import os
import datetime
import base64
from io import BytesIO
from docx import Document

# ËÆæÁΩÆÈ°µÈù¢ÈÖçÁΩÆ
st.set_page_config(
    page_title="My Fan World",
    page_icon="üé®",
    layout="wide"
)

# Ê∑ªÂä†Ëá™ÂÆö‰πâCSS
st.markdown("""
    <style>
        /* ÈöêËóèÈ°µÈù¢È°∂Á´ØÁöÑÁôΩËâ≤ÈïøÊù° */
        #MainMenu {
            visibility: hidden;
        }
        footer {
            visibility: hidden;
        }
        header {
            visibility: hidden;
        }
        .stDeployButton {
            visibility: hidden;
        }
        /* Ë∞ÉÊï¥È°µÈù¢ËæπË∑ù */
        .main .block-container {
            padding-top: 0;
            padding-bottom: 0;
        }
        /* ÂÖ∂‰ªñÊ†∑Âºè */
        .stApp {
            max-width: 1200px;
            margin: 0 auto;
        }
        .title {
            text-align: center;
            color: #262730;
            margin-bottom: 30px;
        }
        .content {
            background-color: #f0f2f6;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }
    </style>
    """, unsafe_allow_html=True)

# Ëá™ÂÆö‰πâCSSÊ†∑Âºè
st.markdown("""
<style>
    .main {
        background-color: #faf0ff;
    }
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .title {
        color: #9c27b0;
        font-size: 2.5em;
        text-align: center;
        margin-bottom: 1em;
    }
    .subtitle {
        color: #673ab7;
        font-size: 1.5em;
        margin-bottom: 1em;
    }
    .content {
        background-color: white;
        padding: 2em;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
</style>
""", unsafe_allow_html=True)

# ÂàùÂßãÂåñ session state
if 'works' not in st.session_state:
    st.session_state.works = {
        "Fan Art": [
            {"title": "OC", "description": "I tried my best...", "date": "2025-05-15", "file": "images/oc.jpg"},
            {"title": "Big Nostril Sheep", "description": "Avatar", "date": "2025-03-10", "file": "images/sheep.jpg"},
            {"title": "Meme", "description": "This is my CP", "date": "2024-03-05", "file": "images/meme.jpg"},
        ],
        "Fan Fiction": [
            {"title": "„ÄêGintoki x Utsuro„ÄëGrowing Pains", "description": "This is my first fan fiction", "date": "2024-03-12", "file": "documents/growing_pains.docx"}
        ],
        "Handmade Works": [
            {"title": "Baby Matthew", "description": "Taking baby to the flower market", "date": "2025-03-05", "file": "images/baby_matthew.jpg"},
            {"title": "Persona 5 Pain Bag", "description": "!!", "date": "2025-02-08", "file": "images/persona5_bag.jpg"},
            {"title": "King Fan", "description": "How is it?", "date": "2024-12-21", "file": "images/king_fan.jpg"},
        ]
    }

if 'collections' not in st.session_state:
    st.session_state.collections = {
        "Fan Art": [
            {
                "title": "Green Frog",
                "author": "Wolftree",
                "source": "Xiaohongshu",
                "link": "https://www.xiaohongshu.com/discovery/item/6834226f000000002100b445?source=webshare&xhsshare=pc_web&xsec_token=CBDPHeSOzae1OP7aitl1Fg384Sw8iNeqd_CdKf__lQVSo=&xsec_source=pc_share",
                "tags": ["#GreenFrog", "#SultanGame"],
                "date": "2025-04-27"
            },
            {
                "title": "„ÄêGintoki x Hijikata„ÄëHuh?!",
                "author": "Xiaohongshu Cat Lover",
                "source": "Xiaohongshu",
                "link": "https://www.xiaohongshu.com/discovery/item/66c856bb000000001d01b280?source=webshare&xhsshare=pc_web&xsec_token=AB7514JeEG0KuMjXy_RYyR4_Ds4jNykcAg3ZDtuIJlmYI=&xsec_source=pc_share",
                "tags": ["#GintokiHijikata", "#Cute", "#Daily"],
                "date": "2025-03-11"
            },
            {
                "title": "„ÄêSougo x Kagura„ÄëSomeone's Being Tsundere Again",
                "author": "LongKissGoodBye",
                "source": "Xiaohongshu",
                "link": "https://www.xiaohongshu.com/discovery/item/682ca244000000000f03a8b7?source=webshare&xhsshare=pc_web&xsec_token=ABSKCD_7_D_4kVceaaywRSm24mR4hnQHbHEw2NrpWEpCA=&xsec_source=pc_share",
                "tags": ["#SougoKagura", "#Gintama", "#FanArt"],
                "date": "2024-12-02"
            }
        ],
        "Fan Fiction": [
            {"title": "Is Being a Teacher's Wife Also a High-Risk Job?", 
             "author": "Phone Computer Freeze", 
             "source": "Jinjiang Literature City", 
             "link": "https://www.xiaohongshu.com/discovery/item/65b4cf6d000000000c0052f0?source=webshare&xhsshare=pc_web&xsec_token=ABHI-k6Z3Uqu-upcQK9O-RDgLqTcjir9P0vdSUn2YKlsA=&xsec_source=pc_share", 
             "tags": ["#Gintama", "#Shoyo", "#Gintoki"], 
             "date": "2025-04-15"},
            {"title": "„ÄêGintoki x Katsura„ÄëWhen Wooden Sword Cuts Stars", 
             "author": "Zero Three Present (School Started)", 
             "source": "Xiaohongshu", 
             "link": "https://www.xiaohongshu.com/discovery/item/679c9445000000001902cabc?source=webshare&xhsshare=pc_web&xsec_token=ABO9ZBWzS61aPegdG0qFNTDEiuFbFPRtRz4cI5_yGIYbU=&xsec_source=pc_share", 
             "tags": ["#Gintama", "#Katsura", "#Gintoki", "#GinKatsu"], 
             "date": "2025-03-10"},
            {"title": "Gintoki and Katsura's Daily Life", 
             "author": "Xiaohongshu Cat Lover", 
             "source": "Xiaohongshu", 
             "link": "https://www.xiaohongshu.com/discovery/item/679c9445000000001902cabc?source=webshare&xhsshare=pc_web&xsec_token=ABO9ZBWzS61aPegdG0qFNTDEiuFbFPRtRz4cI5_yGIYbU=&xsec_source=pc_share", 
             "tags": ["#Gintama", "#GinKatsu"], 
             "date": "2025-03-10"}
        ],
        "Fan Videos": [
            {
                "title": "Artur's At Your Service",
                "author": "Life's Film Suddenly Flashes",
                "source": "Bilibili",
                "link": "https://www.bilibili.com/video/BV1GJ4m1Y7Yd/?spm_id_from=333.999.0.0&vd_source=8a0c7c1c0c0c0c0c0c0c0c0c0c0c0c0c",
                "tags": ["#SultanGame", "#Tusu", "#Artur"],
                "date": "2025-05-26"
            },
            {
                "title": "„ÄêSultan Game„ÄëEmpty Throne",
                "author": "Algae",
                "source": "Bilibili",
                "link": "https://www.bilibili.com/video/BV1GJ4m1Y7Yd/?spm_id_from=333.999.0.0&vd_source=8a0c7c1c0c0c0c0c0c0c0c0c0c0c0c0c",
                "tags": ["#SultanGame", "#HandDrawn", "#EmptyThrone"],
                "date": "2024-03-26"
            }
        ]
    }

if 'logs' not in st.session_state:
    st.session_state.logs = [
        {"title": "Creative Inspiration", "content": "Got new creative inspiration today...", "date": "2025-03-08", "type": "Creative Thoughts"},
        {"title": "Work Completed", "content": "Finally finished the new work...", "date": "2025-02-10", "type": "Creative Process"}
    ]

if 'current_page' not in st.session_state:
    st.session_state.current_page = "üè† Home"

# ‰æßËæπÊ†èÂØºËà™
st.sidebar.title("Navigation Menu")
page = st.sidebar.radio(
    "Select Page",
    ["üè† Home", "üìö Works", "‚ù§Ô∏è Collections", "üìù Creative Log", "üë§ About Me"],
    key="page",
    on_change=lambda: setattr(st.session_state, 'current_page', st.session_state.page)
)

# ‰∏ªÈ°µ
if st.session_state.current_page == "üè† Home":
    st.markdown('<h1 class="title">Welcome to My Fan World</h1>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown('<div class="content">', unsafe_allow_html=True)
        st.markdown("""
        ### Latest Updates
        - üé® New Work Released: "Growing Pains"
        - üìÖ Creative Log Updated: New Creative Plan
        - üí´ Collections Updated
        
        ### Recent Activities
        - Attended CP31 Fan Convention
        - Completed Watching New Works
        - Updated Personal Website Information
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="content">', unsafe_allow_html=True)
        st.markdown("### Quick Links")
        
        if st.button("üìö Works", key="to_works"):
            st.session_state.current_page = "üìö Works"
            st.rerun()
            
        if st.button("‚ù§Ô∏è Collections", key="to_collections"):
            st.session_state.current_page = "‚ù§Ô∏è Collections"
            st.rerun()
            
        if st.button("üìù Creative Log", key="to_logs"):
            st.session_state.current_page = "üìù Creative Log"
            st.rerun()
            
        if st.button("üë§ About Me", key="to_about"):
            st.session_state.current_page = "üë§ About Me"
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

# ‰ΩúÂìÅÂ±ïÁ§∫
elif st.session_state.current_page == "üìö Works":
    if st.button("üè† Back to Home", key="back_to_home_works"):
        st.session_state.current_page = "üè† Home"
        st.rerun()
    
    st.markdown('<h1 class="title">My Works</h1>', unsafe_allow_html=True)
    
    with st.expander("Add New Work", expanded=False):
        with st.form("new_work"):
            st.subheader("Upload New Work")
            work_type = st.selectbox("Work Type", ["Fan Art", "Fan Fiction", "Handmade Works"])
            title = st.text_input("Title")
            description = st.text_area("Description")
            date = st.date_input("Creation Date")
            
            uploaded_file = None
            if work_type in ["Fan Art", "Handmade Works"]:
                uploaded_file = st.file_uploader("Upload Image", type=['png', 'jpg', 'jpeg'])
                if uploaded_file is not None:
                    st.image(uploaded_file, caption="Preview", width=300)
            elif work_type == "Fan Fiction":
                uploaded_file = st.file_uploader("Upload Document", type=['txt', 'doc', 'docx'])
                if uploaded_file is not None:
                    if uploaded_file.type == "text/plain":
                        content = uploaded_file.getvalue().decode()
                        st.text_area("Document Preview", content, height=200)
                    elif uploaded_file.type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                        temp_path = f"temp_{uploaded_file.name}"
                        with open(temp_path, "wb") as f:
                            f.write(uploaded_file.getvalue())
                        try:
                            doc = Document(temp_path)
                            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            st.text_area("Document Preview", content, height=200)
                        except Exception as e:
                            st.error(f"Cannot read Word file: {str(e)}")
                        finally:
                            if os.path.exists(temp_path):
                                os.remove(temp_path)
                    else:
                        st.info("Document preview only supports txt and Word formats")
            
            if st.form_submit_button("Publish Work"):
                if uploaded_file is not None:
                    new_work = {
                        "title": title,
                        "description": description,
                        "date": date.strftime("%Y-%m-%d"),
                        "file": uploaded_file
                    }
                    
                    if work_type in st.session_state.works:
                        st.session_state.works[work_type].append(new_work)
                        st.success("Work published successfully!")
                    else:
                        st.error("Invalid work type!")
                else:
                    st.error("Please upload a file!")

    # ÊòæÁ§∫‰ΩúÂìÅ
    categories = ["All"] + list(st.session_state.works.keys())
    selected_category = st.selectbox("Select Category", categories)
    
    if selected_category == "All":
        for category, items in st.session_state.works.items():
            st.markdown(f"### {category}")
            for i in range(0, len(items), 2):
                cols = st.columns(2)
                for j in range(2):
                    if i + j < len(items):
                        with cols[j]:
                            work = items[i + j]
                            st.markdown(f"""
                            <div style='
                                background-color: white;
                                padding: 15px;
                                border-radius: 10px;
                                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                                margin: 10px;
                                height: 100%;
                            '>
                                <h5 style='margin-bottom: 10px;'>{work['title']}</h5>
                                <p style='margin: 5px 0; color: #666;'>{work['description']}</p>
                                <p style='margin: 5px 0; color: #666;'>Creation Date: {work['date']}</p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            if hasattr(work['file'], 'type'):
                                if work['file'].type.startswith('image/'):
                                    st.image(work['file'], caption="Work Preview", width=300)
                                elif work['file'].type == "text/plain":
                                    content = work['file'].getvalue().decode()
                                    st.text_area("Document Preview", content, height=200)
                                elif work['file'].type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                                    temp_path = f"temp_{work['file'].name}"
                                    with open(temp_path, "wb") as f:
                                        f.write(work['file'].getvalue())
                                    try:
                                        doc = Document(temp_path)
                                        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                                        st.text_area("Document Preview", content, height=200)
                                    except Exception as e:
                                        st.error(f"Cannot read Word file: {str(e)}")
                                    finally:
                                        if os.path.exists(temp_path):
                                            os.remove(temp_path)

# Êî∂ËóèÂ§π
elif st.session_state.current_page == "‚ù§Ô∏è Collections":
    if st.button("üè† Back to Home", key="back_to_home_collections"):
        st.session_state.current_page = "üè† Home"
        st.rerun()
    
    st.markdown('<h1 class="title">My Collections</h1>', unsafe_allow_html=True)
    
    with st.expander("Add New Collection", expanded=False):
        with st.form("new_collection"):
            st.subheader("Add New Collection")
            collection_title = st.text_input("Title")
            collection_author = st.text_input("Author")
            collection_source = st.text_input("Source")
            collection_link = st.text_input("Link")
            collection_tags = st.text_input("Tags (separated by spaces)")
            collection_date = st.date_input("Collection Date")
            collection_category = st.selectbox("Category", ["Fan Art", "Fan Fiction", "Fan Videos"])
            
            if st.form_submit_button("Add Collection"):
                if collection_title and collection_author and collection_source and collection_link:
                    new_collection = {
                        "title": collection_title,
                        "author": collection_author,
                        "source": collection_source,
                        "link": collection_link,
                        "tags": collection_tags.split(),
                        "date": collection_date.strftime("%Y-%m-%d")
                    }
                    st.session_state.collections[collection_category].append(new_collection)
                    st.success("Collection added successfully!")
                else:
                    st.error("Please fill in all required fields!")
    
    selected_category = st.selectbox(
        "Select Category",
        ["All", "Fan Art", "Fan Fiction", "Fan Videos"]
    )
        
    if selected_category == "All":
        for category, items in st.session_state.collections.items():
            st.markdown(f"#### {category}")
            for i in range(0, len(items), 3):
                cols = st.columns(3)
                for j in range(3):
                    if i + j < len(items):
                        with cols[j]:
                            item = items[i + j]
                            st.markdown(f"""
                            <div style='
                                background-color: white;
                                padding: 15px;
                                border-radius: 10px;
                                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                                margin: 10px;
                                height: 100%;
                            '>
                                <h5 style='margin-bottom: 10px;'>{item['title']}</h5>
                                <p style='margin: 5px 0;'>Author: {item['author']}</p>
                                <p style='margin: 5px 0;'>Source: {item['source']}</p>
                                <p style='margin: 5px 0;'>Tags: {' '.join(item['tags'])}</p>
                                <p style='margin: 5px 0;'>Collection Date: {item['date']}</p>
                                <a href="{item['link']}" target="_blank" style='
                                    display: inline-block;
                                    padding: 5px 10px;
                                    background-color: #f0f2f6;
                                    color: #262730;
                                    text-decoration: none;
                                    border-radius: 5px;
                                    margin-top: 10px;
                                '>View Details</a>
                            </div>
                            """, unsafe_allow_html=True)
    elif selected_category in st.session_state.collections:
        items = st.session_state.collections[selected_category]
        for i in range(0, len(items), 3):
            cols = st.columns(3)
            for j in range(3):
                if i + j < len(items):
                    with cols[j]:
                        item = items[i + j]
                        st.markdown(f"""
                        <div style='
                            background-color: white;
                            padding: 15px;
                            border-radius: 10px;
                            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                            margin: 10px;
                            height: 100%;
                        '>
                            <h5 style='margin-bottom: 10px;'>{item['title']}</h5>
                            <p style='margin: 5px 0;'>Author: {item['author']}</p>
                            <p style='margin: 5px 0;'>Source: {item['source']}</p>
                            <p style='margin: 5px 0;'>Tags: {' '.join(item['tags'])}</p>
                            <p style='margin: 5px 0;'>Collection Date: {item['date']}</p>
                            <a href="{item['link']}" target="_blank" style='
                                display: inline-block;
                                padding: 5px 10px;
                                background-color: #f0f2f6;
                                color: #262730;
                                text-decoration: none;
                                border-radius: 5px;
                                margin-top: 10px;
                            '>View Details</a>
                        </div>
                        """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

# Âàõ‰ΩúÊó•Âøó
elif st.session_state.current_page == "üìù Creative Log":
    if st.button("üè† Back to Home", key="back_to_home_logs"):
        st.session_state.current_page = "üè† Home"
        st.rerun()
    
    st.markdown('<h1 class="title">Creative Log</h1>', unsafe_allow_html=True)
        
    with st.expander("Add New Log", expanded=False):
        with st.form("new_log"):
            st.subheader("Add New Log")
            log_title = st.text_input("Title")
            log_content = st.text_area("Content")
            log_date = st.date_input("Date")
            log_type = st.selectbox("Type", ["Creative Plan", "Creative Process", "Creative Thoughts", "Other"])
            
            if st.form_submit_button("Add Log"):
                if log_title and log_content:
                    new_log = {
                        "title": log_title,
                        "content": log_content,
                        "date": log_date.strftime("%Y-%m-%d"),
                        "type": log_type
                    }
                    st.session_state.logs.append(new_log)
                    st.success("Log added successfully!")
                else:
                    st.error("Please fill in the title and content!")
    
    for log in sorted(st.session_state.logs, key=lambda x: x['date'], reverse=True):
        st.markdown(f"""
        <div style='
            background-color: white;
            padding: 15px;
            border-radius: 10px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin: 10px 0;
        '>
            <h4>{log['title']}</h4>
            <p style='color: #666;'>Type: {log['type']} | Date: {log['date']}</p>
            <p>{log['content']}</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

# ÂÖ≥‰∫éÊàë
elif st.session_state.current_page == "üë§ About Me":
    if st.button("üè† Back to Home", key="back_to_home_about"):
        st.session_state.current_page = "üè† Home"
        st.rerun()
    
    st.markdown('<h1 class="title">About Me</h1>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown('<div class="content">', unsafe_allow_html=True)
        st.markdown("### Personal Avatar")
        # ‰ΩøÁî®Âç†‰ΩçÂõæÁâá
        st.image("https://via.placeholder.com/200", width=125)
        
        st.markdown("""
        ### Basic Information
        - Full Name: Chen Leiyu
        - Nickname/CN: Algae
        - MBTI: INTJ
        - Birthday: 2002-11-25
        - Location: China
        
        ### Contact Information
        - Weibo: @Who Calls Me Little Algae
        - Email: chenleiyu1125@foxmail.com
        - QQ: 1078984367
        - WeChat: bzdsmid1125
        """)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="content">', unsafe_allow_html=True)
        st.markdown("""
        ### Composition Table
        - Fated: Gintoki and Hijikata
        - Main CP: Kaguya and Gintoki, Shoyo and Katsura
        - Wallpapers: Artur, Hijikata, Sougo, Katsura
        
        ### Creative Preferences
        - Good at: QQ people, female avatar, short fan fiction
        - Style: Healing, Daily
        - Common Software: PS, AI, SAI, Stone Document
        
        ### Favorite Works
        - Animation: "Gintama", "Naruto", "Jojo's Bizarre Adventure", "Fate stay night"
        - Comic: "Beautiful People - Outside Story", "Night by the Water"
        - Game: "Devil May Cry 5", "Ace Attorney", "Zelda", "Stardew Valley", "Sultan Game"
        - Novel: "Dragon Races", "Human Scum Self-Help System", "Mystery Lord", "Full-Time Master", "No Head Knight's Strange Adventure"

        ### Dislikes
        - Game: "Genshin Impact", "Fifth Person", "Black Myth - Sun Wukong"
        - Animation: "My Hero Academia", "Curse Return Battle", "Wenhao Wild Dog"
        - CP: Bo Junyi and Xiaoxiao, Kekai, Five Summer, Tai Zhong Tai
        """)
        st.markdown('</div>', unsafe_allow_html=True)

# È°µËÑö
st.markdown("""
---
<div style='text-align: center'>
    <p>¬© 2025 Algae's Kingdom | Made with ‚ù§Ô∏è</p>
</div>
""", unsafe_allow_html=True) 